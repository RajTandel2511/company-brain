"""Document intelligence — extract text from files and map entities to Spectrum.

Phase 1: PDF text extraction + regex/fuzzy matching against master lists
(jobs, POs, vendors, customers, AP/AR invoice numbers). Results persist in the
same nas-index.sqlite DB.

Phase 2 (later): OCR for scanned PDFs/images, Office docs, classifier.

The extractor is incremental: only files not yet processed or newer than their
last extraction are touched, so you can run it any time and it picks up.
"""
from __future__ import annotations

import os
import re
import shutil
import sqlite3
import time
from pathlib import Path
from typing import Iterable

from . import nas_index
from .config import settings
from .db import run_query

# --- storage ------------------------------------------------------------

# Separate DB so docintel writes never contend with the NAS indexer's locks.
# We look up file_id -> path via nas_index.DB when we need to open a file.
DB = nas_index.DATA_DIR / "docintel.sqlite"


# Surrogate codepoints (\ud800-\udfff) and NUL (\x00) are valid in Python str
# but cannot be encoded as UTF-8 / stored in SQLite. PDFs with broken font
# encodings sometimes produce them; strip before insert.
_SQL_BAD_CHARS = re.compile(r"[\ud800-\udfff\x00]")


def _connect() -> sqlite3.Connection:
    """Docintel owns its own DB — no cross-attach to nas-index (which would
    contend with the live indexer). File metadata is mirrored into a local
    files table during extraction runs."""
    c = sqlite3.connect(DB, timeout=30)
    c.execute("PRAGMA journal_mode=WAL")
    c.execute("PRAGMA busy_timeout=30000")
    c.executescript("""
        CREATE TABLE IF NOT EXISTS files (
          id INTEGER PRIMARY KEY,
          path TEXT UNIQUE NOT NULL,
          name TEXT NOT NULL,
          size INTEGER,
          modified REAL
        );
        CREATE TABLE IF NOT EXISTS file_content (
          file_id INTEGER PRIMARY KEY, text TEXT,
          char_count INTEGER, extracted_at REAL, extractor TEXT
        );
        CREATE TABLE IF NOT EXISTS file_entities (
          file_id INTEGER NOT NULL, entity_type TEXT NOT NULL,
          entity_value TEXT NOT NULL, source TEXT,
          PRIMARY KEY (file_id, entity_type, entity_value)
        ) WITHOUT ROWID;
        CREATE INDEX IF NOT EXISTS file_entities_lookup
          ON file_entities(entity_type, entity_value);
    """)
    return c


def _nas_conn() -> sqlite3.Connection:
    """Read-only snapshot connection to the NAS index. Short-lived — open,
    query, close — so we never hold a lock while the indexer is writing."""
    c = sqlite3.connect(nas_index.DB, timeout=60)
    c.execute("PRAGMA busy_timeout=60000")
    return c


# --- master lists from Spectrum ----------------------------------------

_YEAR_LIKE = re.compile(r"^(19|20)\d{2}$")
_DATE_LIKE = re.compile(r"^\d{1,2}[./]\d{1,2}[./]\d{2,4}$")
_PLACEHOLDER = {
    "XXXX", "0000", "1111", "9999", "NNNN", "????",
    "TEST", "DEMO", "NULL", "NONE", "SAMPLE", "N/A", "VOID", "CANCEL",
}


def _invoice_match_quality(inv: str) -> str:
    """Classify an invoice number for matching strategy.

    'strong'  - unique enough to match plain substring (alphanumeric or long)
    'weak'    - numeric/short, needs invoice-context nearby to trust
    'skip'    - always ignore (year, date-fragment, placeholder, too short)
    """
    s = inv.strip().upper()
    if len(s) < 4 or s in _PLACEHOLDER:
        return "skip"
    if _YEAR_LIKE.match(s):
        return "skip"
    if _DATE_LIKE.match(s):
        return "skip"
    # "Strong" = contains a letter or is >=7 digits (very unlikely collision)
    if any(c.isalpha() for c in s) or (s.isdigit() and len(s) >= 7):
        return "strong"
    return "weak"


def _load_masters() -> dict:
    # Allow running on environments without Spectrum connectivity (e.g. a
    # cloud VM used for a one-time extraction burst). Entity tagging is
    # skipped — the text content is still extracted, and entity matching
    # can be re-run later by a machine that can reach Spectrum.
    if os.environ.get("SKIP_SPECTRUM") == "1":
        return {
            "jobs": set(), "vendors": {}, "customers": {},
            "pos": set(), "pos_strong": set(), "pos_weak": set(),
            "ap_strong": set(), "ap_weak": set(), "ap_invoices": set(),
        }

    CO = settings.spectrum_company_code
    jobs = {r["v"].strip() for r in run_query(
        f"SELECT LTRIM(RTRIM(Job_Number)) AS v FROM dbo.JC_JOB_MASTER_MC WHERE Company_Code='{CO}'"
    )["rows"] if r.get("v")}

    vendors = {}  # code -> name
    for r in run_query(
        f"SELECT LTRIM(RTRIM(Vendor_Code)) AS c, LTRIM(RTRIM(Vendor_Name)) AS n "
        f"FROM dbo.VN_VENDOR_MASTER_MC WHERE Company_Code='{CO}'"
    )["rows"]:
        code = (r.get("c") or "").upper()
        if len(code) >= 3:
            vendors[code] = r.get("n") or code

    customers = {}
    for r in run_query(
        f"SELECT LTRIM(RTRIM(Customer_Code)) AS c, LTRIM(RTRIM(Name)) AS n "
        f"FROM dbo.CR_CUSTOMER_MASTER_MC WHERE Company_Code='{CO}'"
    )["rows"]:
        code = (r.get("c") or "").upper()
        if len(code) >= 3:
            customers[code] = r.get("n") or code

    pos_raw = {r["v"].strip() for r in run_query(
        f"SELECT LTRIM(RTRIM(PO_Number)) AS v FROM dbo.PO_PURCHASE_ORDER_HEADER_MC WHERE Company_Code='{CO}'"
    )["rows"] if r.get("v")}
    # Same strong/weak bucketing as invoices: unique strings (alpha or 7+ digits)
    # get plain match; 4-6 digit numerics require invoice/PO-context nearby.
    pos_strong, pos_weak = set(), set()
    for po in pos_raw:
        q = _invoice_match_quality(po)  # same rules work for PO numbers
        if q == "strong":
            pos_strong.add(po)
        elif q == "weak":
            pos_weak.add(po)

    # Bucket AP invoices into strong/weak/skip so we can match each appropriately
    ap_raw = {r["v"].strip() for r in run_query(
        f"SELECT DISTINCT LTRIM(RTRIM(Invoice_Number)) AS v FROM dbo.VN_GL_DISTRIBUTION_HEADER_MC "
        f"WHERE Company_Code='{CO}' AND Check_Date >= DATEADD(year, -5, GETDATE())"
    )["rows"] if r.get("v")}
    ap_strong: set[str] = set()
    ap_weak: set[str] = set()
    for inv in ap_raw:
        q = _invoice_match_quality(inv)
        if q == "strong":
            ap_strong.add(inv)
        elif q == "weak":
            ap_weak.add(inv)

    return {
        "jobs": jobs, "vendors": vendors, "customers": customers,
        "pos": pos_raw,  # kept for backwards reference
        "pos_strong": pos_strong,
        "pos_weak": pos_weak,
        "ap_strong": ap_strong,
        "ap_weak": ap_weak,
        "ap_invoices": ap_raw,  # backwards reference only
    }


# --- text extraction ---------------------------------------------------

_JOB_PAT = re.compile(r"(?<!\d)(\d{1,2}\.\d{2,3})(?!\d)")


OCR_MIN_TEXT_CHARS = 80  # below this, assume scanned and fall back to OCR
OCR_MAX_PAGES = 5         # OCR is slow — cap per-PDF page count


def extract_pdf_text(path: Path, max_pages: int = 25) -> tuple[str, str]:
    """Extract text from a PDF.

    Returns (text, extractor_name). Uses pypdf first; if the text layer is
    basically empty, falls back to Tesseract OCR via PyMuPDF rasterization.
    """
    try:
        import pypdf
    except ImportError:
        return "", "none"

    text_layer = ""
    try:
        reader = pypdf.PdfReader(str(path))
        parts = []
        for i, page in enumerate(reader.pages):
            if i >= max_pages:
                break
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        text_layer = "\n".join(parts)
    except Exception:
        text_layer = ""

    if len(text_layer.strip()) >= OCR_MIN_TEXT_CHARS:
        return text_layer, "pypdf"

    # Fall back to OCR
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
        import io
    except ImportError:
        return text_layer, "pypdf"

    try:
        doc = fitz.open(str(path))
    except Exception:
        return text_layer, "pypdf"

    ocr_parts: list[str] = []
    try:
        for i, page in enumerate(doc):
            if i >= OCR_MAX_PAGES:
                break
            try:
                # 200 dpi is a good balance of speed vs accuracy for scanned forms
                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                t = pytesseract.image_to_string(img, lang="eng", config="--psm 6")
                if t:
                    ocr_parts.append(t)
            except Exception:
                continue
    finally:
        doc.close()

    ocr_text = "\n".join(ocr_parts)
    # Prefer whichever is longer (sometimes pypdf gets partial text we'd lose)
    return (ocr_text if len(ocr_text) > len(text_layer) else text_layer), "ocr"


def extract_xlsx_text(path: Path, max_sheets: int = 10, max_cells: int = 20_000) -> str:
    """Extract visible cell text from an .xlsx workbook."""
    try:
        import openpyxl
    except ImportError:
        return ""
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return ""
    parts: list[str] = []
    cells_read = 0
    for sname in wb.sheetnames[:max_sheets]:
        ws = wb[sname]
        parts.append(f"\n=== SHEET {sname} ===\n")
        for row in ws.iter_rows(values_only=True):
            for v in row:
                if v is None:
                    continue
                parts.append(str(v))
                cells_read += 1
                if cells_read >= max_cells:
                    break
            if cells_read >= max_cells:
                break
        if cells_read >= max_cells:
            break
    wb.close()
    return " ".join(parts)


def extract_docx_text(path: Path) -> str:
    """Extract paragraph + table text from a .docx."""
    try:
        import docx
    except ImportError:
        return ""
    try:
        d = docx.Document(str(path))
    except Exception:
        return ""
    parts = [p.text for p in d.paragraphs if p.text]
    for tbl in d.tables:
        for row in tbl.rows:
            parts.extend(c.text for c in row.cells if c.text)
    return "\n".join(parts)


def _pdf_wrapper(path: Path) -> str:
    """Adapter: PDF extractor returns (text, extractor). Drops extractor name."""
    text, _ = extract_pdf_text(path)
    return text


EXTRACTORS = {
    ".pdf":  ("pypdf+ocr", _pdf_wrapper),
    ".xlsx": ("openpyxl",  extract_xlsx_text),
    ".xlsm": ("openpyxl",  extract_xlsx_text),
    ".docx": ("python-docx", extract_docx_text),
}


def entities_from_text(text: str, masters: dict) -> list[tuple[str, str, str]]:
    """Return (entity_type, entity_value, source) triples extracted from text."""
    if not text:
        return []
    out: set[tuple[str, str, str]] = set()
    up = text.upper()

    # Jobs — match regex then verify against master list
    for m in _JOB_PAT.finditer(text):
        v = m.group(1)
        if v in masters["jobs"]:
            out.add(("job", v, "regex+match"))

    # POs — strong bucket gets plain substring (unique enough);
    #       weak bucket (short numeric) needs "PO"/"P.O."/"order #" context nearby.
    for po in masters.get("pos_strong", ()):
        if po and po in text:
            out.add(("po", po, "strong"))

    weak_pos = masters.get("pos_weak", ())
    if weak_pos:
        lower = text.lower()
        po_keyword_positions: list[int] = []
        for kw in ("po #", "po#", "po:", "p.o.", "p.o #", "purchase order", "order #", "order#"):
            start = 0
            while True:
                i = lower.find(kw, start)
                if i < 0: break
                po_keyword_positions.append(i)
                start = i + 1
        if po_keyword_positions:
            for m in re.finditer(r"\b(\d{4,10})\b", text):
                v = m.group(1)
                if v not in weak_pos:
                    continue
                pos2 = m.start()
                if any(abs(pos2 - kp) <= 40 for kp in po_keyword_positions):
                    out.add(("po", v, "context"))

    # Vendor / customer codes
    tokens = set(re.findall(r"[A-Z0-9]{3,10}", up))
    for t in tokens:
        if t in masters["vendors"]:
            out.add(("vendor", t, "code"))
        if t in masters["customers"]:
            out.add(("customer", t, "code"))

    # Vendor / customer names (case-insensitive substring of normalized name)
    for code, name in masters["vendors"].items():
        if name and len(name) >= 6 and name.upper() in up:
            out.add(("vendor", code, "name"))
    for code, name in masters["customers"].items():
        if name and len(name) >= 6 and name.upper() in up:
            out.add(("customer", code, "name"))

    # AP invoices — strict matching to avoid year / date / short-number collisions
    # 1) Strong invoices (alphanumeric or 7+ digits): plain substring — very low collision risk
    for inv in masters.get("ap_strong", ()):
        if inv and inv in text:
            out.add(("ap_invoice", inv, "strong"))

    # 2) Weak invoices (short numeric / ambiguous): require "invoice"/"inv"/"bill"
    #    keyword within 40 characters of the match, AND word-boundary enforcement.
    weak = masters.get("ap_weak", ())
    if weak:
        # Extract candidate numeric tokens from the text that are near an invoice-context cue
        lower = text.lower()
        keyword_positions: list[int] = []
        for kw in ("invoice", "inv#", "inv #", "inv. ", "inv:", "bill #", "bill#", "bill no", "ref #", "ref#"):
            start = 0
            while True:
                i = lower.find(kw, start)
                if i < 0: break
                keyword_positions.append(i)
                start = i + 1
        if keyword_positions:
            for m in re.finditer(r"\b(\d{4,10})\b", text):
                v = m.group(1)
                if v not in weak:
                    continue
                pos = m.start()
                # Is there an invoice-related keyword within 40 chars on either side?
                if any(abs(pos - kp) <= 40 for kp in keyword_positions):
                    out.add(("ap_invoice", v, "context"))

    return list(out)


# --- worker ------------------------------------------------------------

SUPPORTED_EXT_SQL = " OR ".join(
    f"lower(f.name) LIKE '%{ext}'" for ext in EXTRACTORS
)


def _extract_one(full_path: Path, ext: str) -> tuple[str, str]:
    """Worker-thread function: extract text, return (text, extractor_name)."""
    info = EXTRACTORS.get(ext)
    if not info:
        return "", "unsupported"
    extractor_name, extractor_fn = info
    if ext == ".pdf":
        return extract_pdf_text(full_path)  # returns (text, used)
    try:
        return extractor_fn(full_path), extractor_name
    except Exception:
        return "", "failed"


def run(batch_limit: int = 500, progress_cb=None, workers: int | None = None) -> dict:
    """Process up to `batch_limit` unprocessed files and extract + link.

    Extraction runs in parallel (thread pool). SQLite writes stay serial on the
    main thread. Supported: PDF (pypdf → OCR fallback), XLSX/XLSM, DOCX.
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed
    import os as _os

    if workers is None:
        workers = min(8, max(2, (_os.cpu_count() or 4) // 2))

    # When the NAS is mounted over a slow link (Tailscale-SMB on the GCP burst
    # node), per-file SMB latency dominates extraction. EXTRACT_PREFETCH=1 turns
    # on a bulk-copy step so we read each batch off SMB once into a local SSD
    # cache, then extract from local — pure CPU, no network waiting.
    prefetch = _os.environ.get("EXTRACT_PREFETCH") == "1"
    prefetch_workers = int(_os.environ.get("EXTRACT_PREFETCH_WORKERS", "32"))
    cache_dir = Path(_os.environ.get("EXTRACT_CACHE_DIR", "/tmp/extract-cache"))

    t0 = time.time()
    nas_index.ensure_built()
    masters = _load_masters()

    c = _connect()

    # 1) Pull candidate file metadata from the NAS index
    nc = _nas_conn()
    nas_rows = nc.execute(f"""
        SELECT f.id, f.path, f.modified, f.name, f.size
        FROM files f
        WHERE f.is_dir = 0
          AND ({SUPPORTED_EXT_SQL})
        ORDER BY f.modified DESC
    """).fetchall()
    nc.close()

    # 2) Filter out files we've already extracted
    already = dict(c.execute(
        "SELECT file_id, extracted_at FROM file_content"
    ).fetchall())
    candidates = []
    for file_id, path, mtime, name, size in nas_rows:
        prev = already.get(file_id)
        if prev is not None and prev >= (mtime or 0):
            continue
        c.execute(
            "INSERT OR REPLACE INTO files(id,path,name,size,modified) VALUES (?,?,?,?,?)",
            (file_id, path, name, size, mtime),
        )
        candidates.append((file_id, path, mtime, (name or "").lower()))
        if len(candidates) >= batch_limit:
            break
    c.commit()

    processed = 0
    entity_count = 0
    failed = 0
    root = Path(settings.nas_root)

    # Bulk-prefetch the batch from NAS to local SSD if asked. Done in one pass
    # with a wide thread pool so we saturate the network instead of trickling
    # one file per worker thread during extraction.
    cache_lookup: dict[int, Path] = {}
    if prefetch and candidates:
        cache_dir.mkdir(parents=True, exist_ok=True)
        prefetch_t0 = time.time()

        def _stage(args):
            file_id, rel_path, _mtime, lname = args
            ext = Path(lname).suffix
            src = root / rel_path
            dst = cache_dir / f"{file_id}{ext}"
            try:
                shutil.copy2(src, dst)
                return file_id, dst
            except Exception:
                return file_id, None

        with ThreadPoolExecutor(max_workers=prefetch_workers) as pool:
            for fid, path in pool.map(_stage, candidates):
                if path is not None:
                    cache_lookup[fid] = path
        elapsed = time.time() - prefetch_t0
        hits = len(cache_lookup)
        print(f"  prefetched {hits}/{len(candidates)} files in {elapsed:.1f}s", flush=True)

    def _task(args):
        file_id, rel_path, mtime, lname = args
        ext = Path(lname).suffix
        full = cache_lookup.get(file_id) or (root / rel_path)
        try:
            text, used = _extract_one(full, ext)
        except Exception:
            return (file_id, None, "failed")
        ents = entities_from_text(text, masters)
        return (file_id, (text, used, ents), None)

    cur = c.cursor()
    cur.execute("BEGIN")
    with ThreadPoolExecutor(max_workers=workers) as pool:
        futures = [pool.submit(_task, item) for item in candidates]
        for fut in as_completed(futures):
            file_id, result, err = fut.result()
            if err or result is None:
                failed += 1
                continue
            text, used, ents = result
            # PDFs with broken font CMaps occasionally produce text containing
            # lone UTF-16 surrogate codepoints. SQLite's UTF-8 encoder can't
            # store them and raises UnicodeEncodeError, which used to kill
            # the entire batch transaction (and the worker process). Strip
            # them — and NULs while we're at it — before insert.
            safe_text = _SQL_BAD_CHARS.sub("", text)[:50000] if text else None
            try:
                cur.execute(
                    "INSERT OR REPLACE INTO file_content(file_id,text,char_count,extracted_at,extractor) "
                    "VALUES (?,?,?,?,?)",
                    (file_id, safe_text, len(text or ""), time.time(), used),
                )
                cur.execute("DELETE FROM file_entities WHERE file_id = ?", (file_id,))
                for etype, eval, src in ents:
                    cur.execute(
                        "INSERT OR IGNORE INTO file_entities(file_id,entity_type,entity_value,source) "
                        "VALUES (?,?,?,?)",
                        (file_id, etype, eval, src),
                    )
                    entity_count += 1
            except (UnicodeEncodeError, sqlite3.DataError) as e:
                # Belt-and-suspenders: if cleaning missed something, drop the
                # one file and keep the batch going instead of aborting.
                print(f"  skip file_id={file_id} ({type(e).__name__}: {str(e)[:80]})", flush=True)
                failed += 1
                continue
            processed += 1
            if processed % 20 == 0:
                cur.execute("COMMIT")
                # Yield the write lock so the RAG worker can slip in. Without
                # this pause, COMMIT → BEGIN is microseconds apart and RAG
                # starves out — busy_timeout doesn't help when the lock is
                # held ~continuously.
                time.sleep(0.5)
                cur.execute("BEGIN")
                if progress_cb:
                    progress_cb(processed, entity_count)
    cur.execute("COMMIT")

    # Reclaim the local cache so disk usage stays bounded. Done after COMMIT
    # so a crash mid-batch can resume by re-fetching the same files (cheap).
    if cache_lookup:
        for path in cache_lookup.values():
            try:
                path.unlink(missing_ok=True)
            except Exception:
                pass

    # Update a simple meta marker for UI
    c.execute("""
        CREATE TABLE IF NOT EXISTS meta (k TEXT PRIMARY KEY, v TEXT)
    """)
    c.execute("REPLACE INTO meta(k,v) VALUES ('docintel_last_run', ?)",
              (str(int(time.time())),))
    c.commit()
    c.close()

    return {
        "processed": processed,
        "entities_extracted": entity_count,
        "failed": failed,
        "workers": workers,
        "batch_limit": batch_limit,
        "elapsed_seconds": round(time.time() - t0, 1),
    }


# --- query -------------------------------------------------------------

def files_for_entity(entity_type: str, entity_value: str, limit: int = 200) -> list[dict]:
    c = _connect()
    rows = c.execute("""
        SELECT f.id, f.path, f.name, f.size, f.modified, fe.source
        FROM file_entities fe
        JOIN files f ON f.id = fe.file_id
        WHERE fe.entity_type = ? AND fe.entity_value = ?
        ORDER BY f.modified DESC
        LIMIT ?
    """, (entity_type, entity_value.strip(), limit)).fetchall()
    c.close()
    return [
        {"file_id": r[0], "path": r[1], "name": r[2], "size": r[3], "modified": r[4], "match": r[5]}
        for r in rows
    ]


def stats() -> dict:
    c = _connect()
    n_ext = c.execute("SELECT COUNT(*) FROM file_content").fetchone()[0]
    n_ents = c.execute("SELECT COUNT(*) FROM file_entities").fetchone()[0]
    by = dict(c.execute(
        "SELECT entity_type, COUNT(*) FROM file_entities GROUP BY entity_type"
    ).fetchall())
    # Pending = files that exist in the NAS index but we haven't extracted yet.
    # Short-lived read-only connection to avoid contention.
    try:
        nc = _nas_conn()
        total_supported = nc.execute(f"""
            SELECT COUNT(*) FROM files f
            WHERE f.is_dir = 0 AND ({SUPPORTED_EXT_SQL})
        """).fetchone()[0]
        nc.close()
    except Exception:
        total_supported = n_ext
    pending = max(0, total_supported - n_ext)
    by_extractor = dict(c.execute(
        "SELECT extractor, COUNT(*) FROM file_content GROUP BY extractor"
    ).fetchall())
    # Top entities by linked-file count — great for quick exploration
    top_jobs = c.execute("""
        SELECT entity_value, COUNT(*) AS n
        FROM file_entities WHERE entity_type = 'job'
        GROUP BY entity_value ORDER BY n DESC LIMIT 10
    """).fetchall()
    top_vendors = c.execute("""
        SELECT entity_value, COUNT(*) AS n
        FROM file_entities WHERE entity_type = 'vendor'
        GROUP BY entity_value ORDER BY n DESC LIMIT 10
    """).fetchall()
    c.close()
    return {
        "built": True,
        "files_extracted": n_ext,
        "entities_indexed": n_ents,
        "by_entity_type": by,
        "by_extractor": by_extractor,
        "pending": pending,
        "pdfs_pending": pending,  # backwards-compat for worker script
        "top_jobs": [{"value": r[0], "count": r[1]} for r in top_jobs],
        "top_vendors": [{"value": r[0], "count": r[1]} for r in top_vendors],
    }
